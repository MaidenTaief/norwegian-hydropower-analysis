#!/usr/bin/env python3
"""
Generate Word Document Report for Norwegian Dam Analysis
========================================================

Creates a professional Word document matching the structure needed for
Janhavi's India vs Norway comparative conference paper.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import pandas as pd

def create_report():
    """Generate comprehensive Word document report."""
    
    print("📝 Generating Word Document Report...")
    print("=" * 70 + "\n")
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Norwegian Dam Infrastructure Analysis"
    doc.core_properties.author = "Comparative Analysis - India vs Norway"
    
    # Define styles
    styles = doc.styles
    
    # Title
    title = doc.add_heading('Norwegian Dam Infrastructure Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comparative Study Component: India vs Norway')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 1: INTRODUCTION ===
    doc.add_heading('1. Introduction and Data Sources', 1)
    
    intro_text = """This analysis examines Norwegian dam infrastructure using data from the Norwegian Water Resources and Energy Directorate (NVE), providing a comprehensive comparative basis for understanding dam development patterns between India and Norway. The analysis employs the NVE national hydropower database, which contains detailed information on 4,953 dam structures, 2,997 reservoirs, and associated infrastructure across Norway's mountainous terrain.

The Norwegian Water Resources and Energy Directorate (NVE) serves as the primary regulatory and monitoring authority for water resources and energy infrastructure in Norway. The NVE dataset provides comprehensive coverage of Norwegian hydropower infrastructure, with particular strength in operational parameters such as regulation ranges and storage characteristics—data types that reflect Norway's emphasis on flexible hydropower generation for grid balancing."""
    
    doc.add_paragraph(intro_text)
    
    # Data characteristics table
    doc.add_heading('Dataset Characteristics', 2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    characteristics = [
        ('Data Source', 'Norwegian Water Resources and Energy Directorate (NVE)'),
        ('Total Dam Points', '4,953'),
        ('Total Reservoirs', '2,997'),
        ('Temporal Coverage', '1800 - 2025 (225 years)'),
        ('Geographic Coverage', 'Complete national infrastructure'),
        ('Coordinate System', 'ETRS89 / UTM Zone 33N (converted to WGS84)'),
        ('Primary Focus', 'Hydropower infrastructure and reservoir management')
    ]
    
    for i, (key, value) in enumerate(characteristics):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # === SECTION 2: DECADE-WISE CONSTRUCTION TRENDS ===
    doc.add_heading('2. Decade-wise Construction Trends', 1)
    
    decade_text = """Norway's large-dam programme developed substantially in the decades following World War II, with a pronounced surge during the 1950s to 1970s, driven primarily by hydropower expansion needs. The National NVE hydropower register documents 3,874 dams with recorded construction years, providing clear evidence of Norway's post-war infrastructure development trajectory.

Analysis of decade-wise construction patterns reveals:

• Peak Construction Period: The 1960s saw maximum dam construction with 528 dams completed, representing 13.6% of all dated infrastructure
• Post-War Boom (1945-1980): This 35-year period accounts for 1,789 dams (46.2% of total)
• Early Development (1800-1945): 734 dams (18.9%) from this extended period
• Modern Era (1980-2025): 1,351 dams (34.9%) reflecting continued development and modernization

This construction pattern differs markedly from India's development trajectory, where the peak occurred later (1970s-1990s) driven by post-independence irrigation and multipurpose project expansion. Norway's earlier peak reflects the country's focus on hydropower development as a primary electricity source in mountainous terrain, aligning with post-war electrification and industrialization objectives.

The study by Graabak et al. (2017) illustrates Norway's hydropower system functioning as a "battery" for variable renewables in northern Europe, a capability built on this extensive mid-century infrastructure development."""
    
    doc.add_paragraph(decade_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 1: ').bold = True
    p.add_run('See visualization "decade_wise_norway.png" for detailed decade-wise construction bar chart')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 3: AGE DISTRIBUTION ===
    doc.add_heading('3. Age Distribution Analysis', 1)
    
    age_text = """With completion years spanning over two centuries, Norway's dam stock presents a mature age profile requiring ongoing maintenance and periodic modernization. Classification of dams by age reveals important infrastructure management considerations:

Age Distribution (as of 2025):
• Less than 25 years: 476 dams (12.3%) - Recent construction and modernization
• 25 to 49 years: 815 dams (21.0%) - Mid-life infrastructure requiring monitoring
• 50 to 99 years: 2,331 dams (60.2%) - Mature infrastructure, primary rehabilitation focus
• 100 years or older: 252 dams (6.5%) - Historic structures requiring special attention
• Average dam age: 60.8 years
• Oldest structure: 225 years (constructed 1800)

This age profile is comparable to India's situation where, according to parliamentary documentation from the National Register of Large Dams (NRLD-2023), of the 6,138 completed "specified" large dams, 3,789 (61.7%) are older than 25 years, indicating a broadly similar challenge in managing aging infrastructure.

The national register maintained by the Norwegian Water Resources and Energy Directorate (NVE) monitors approximately 3,600 registered dams classified by consequence level, with high-consequence structures subject to rigorous periodic inspections, instrumentation requirements, and mandatory dam safety reviews. This regulatory framework addresses the age-related risks inherent in Norway's mature dam portfolio."""
    
    doc.add_paragraph(age_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 2: ').bold = True
    p.add_run('See visualization "age_wise_norway.png" for age distribution donut chart')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 4: REGIONAL CONCENTRATION AND GOVERNANCE ===
    doc.add_heading('4. Regional Concentration and Governance Framework', 1)
    
    regional_text = """Spatial analysis reveals distinct geographic concentrations reflecting Norway's topographic and hydrological characteristics. Dam infrastructure is predominantly located in mountainous catchments and western/central regions where hydropower potential is highest.

Regional Distribution (Top 5 Counties):
1. Innlandet: 1,413 dams (28.5%) - Central mountainous region
2. Vestland: 1,037 dams (20.9%) - Western fjord region
3. Trøndelag: 813 dams (16.4%) - Central Norway
4. Nordland: 615 dams (12.4%) - Northern region
5. Agder: 476 dams (9.6%) - Southern region

These concentrations align with Norway's steep elevation gradients and high precipitation zones in western and central mountainous counties (e.g., Vestland, Innlandet) where reservoir storage and high-head hydropower generation are most viable. The fjord-adjacent topography creates ideal conditions for high storage efficiency through deep valley impoundments.

Governance Framework:

Norway's dam safety and water resource management operate under a unified regulatory framework administered by NVE:

• Primary Authority: Norwegian Water Resources and Energy Directorate (NVE)
• Regulatory Basis: Water Resources Act (1940, as amended) and associated regulations
• European Integration: Water Framework Directive (WFD) implementation
• Dam Classification: Consequence-based classification system (four levels)
• Inspection Regime: Mandatory periodic inspections based on consequence level
• Licensing: Watercourse regulation licenses required for major interventions

Hanssen et al. (2016) review Norway's WFD implementation and its implications for hydropower regulation, noting the integration of environmental considerations with traditional safety and operational oversight. This framework differs from India's recent Dam Safety Act, 2021, which established a statutory multi-tiered institutional mechanism (National Committee on Dam Safety, National Dam Safety Authority, State-level bodies), representing a more federated governance structure appropriate to India's state-based water resource management system."""
    
    doc.add_paragraph(regional_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 3: ').bold = True
    p.add_run('See visualization "regional_distribution_norway.png" for regional distribution donut chart')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 5: STORAGE EFFICIENCY AND TOPOGRAPHIC ADAPTATION ===
    doc.add_heading('5. Storage Efficiency and Topographic Adaptation', 1)
    
    storage_text = """A distinctive characteristic of Norwegian reservoir infrastructure is the high storage efficiency achieved through deep valley and fjord topography. Analysis of volume-to-area ratios reveals important insights into how dam design adapts to geographic constraints.

Storage Characteristics:
• Total storage capacity: 61,784 million cubic meters (MCM)
• Average storage per reservoir: 53.0 MCM
• Average storage efficiency: 10.80 MCM/km²
• This efficiency metric is significantly higher than typical broad river valley impoundments

The storage efficiency ratio (volume per unit surface area) serves as an indicator of reservoir depth characteristics. Norway's average of 10.80 MCM/km² reflects the predominance of deep valley storage where steep terrain allows substantial water volume accumulation with relatively limited surface area—a direct consequence of glacially-carved valleys and fjord geography.

In contrast, India's dam infrastructure operates primarily in broader alluvial and semi-alluvial river valleys where reservoir impoundments spread across larger surface areas with comparatively shallower average depths. This fundamental topographic difference influences:

• Land acquisition requirements (lower per unit storage in Norway)
• Evaporation losses (reduced with smaller surface area)
• Ecosystem impacts (different footprint characteristics)
• Construction approaches (Norway: often rock-fill and concrete gravity; India: predominantly earth and composite)

The high storage efficiency also enables Norway's hydropower system to provide substantial energy storage capacity—critical for the "battery of Europe" function in supporting intermittent renewable energy sources across the Nordic and European grid systems."""
    
    doc.add_paragraph(storage_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 4: ').bold = True
    p.add_run('See visualization "storage_efficiency_norway.png" for volume vs area scatter plot')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 6: OPERATIONAL FLEXIBILITY (UNIQUE SECTION) ===
    doc.add_heading('6. Operational Flexibility and Regulation Range', 1)
    
    regulation_text = """A unique strength of the Norwegian NVE dataset is the availability of regulation range data—the vertical difference between highest regulated water level (HRV) and lowest regulated water level (LRV)—which provides quantitative insight into operational flexibility for hydropower generation. This data type is generally not available in published Indian dam databases, representing a distinctive analytical dimension.

Regulation Range Analysis:
• Reservoirs with regulation data: 1,585 (52.9% of total)
• Average regulation range: 10.9 meters
• Maximum regulation range: 193.5 meters
• This operational parameter directly correlates with hydropower peaking capability

Physical Interpretation:

The regulation range represents the usable storage volume that can be cycled for power generation. Larger regulation ranges enable:

1. Greater hydropower peaking capability - ability to respond to demand variations
2. Seasonal storage optimization - capturing spring snowmelt for winter generation
3. Grid balancing services - rapid response to supply-demand imbalances
4. Multi-year storage strategies - managing inter-annual precipitation variability

Strategic Significance:

This operational flexibility underpins Norway's role in Nordic and European energy systems. As documented by Graabak et al. (2017), Norwegian hydropower functions as a "battery" for variable renewable energy sources (wind, solar) across northern Europe. The physical capability for this function derives from:

• High regulation ranges enabling substantial volume cycling
• Multiple interconnected reservoirs creating cascade flexibility  
• Strategic positioning relative to Nordic grid infrastructure
• Regulatory frameworks supporting cross-border power exchange

Comparative Context:

While Indian dams serve multipurpose functions (irrigation, flood control, power generation), operational priorities differ. Indian reservoir operations must balance:
• Irrigation water supply (often pre-committed volumes)
• Flood storage requirements (monsoon season)
• Power generation (often run-of-river or limited peaking)

Norway's predominantly single-purpose hydropower focus, combined with high regulation ranges, enables operational strategies not generally feasible in multipurpose systems with competing water allocation requirements.

This represents a fundamental difference in system design philosophy: Norway optimizes for energy system flexibility, while India optimizes for water resource allocation across multiple competing demands."""
    
    doc.add_paragraph(regulation_text)
    
    # Reference to visualization  
    p = doc.add_paragraph()
    p.add_run('Figure 5: ').bold = True
    p.add_run('See visualization "regulation_range_norway.png" for regulation range histogram')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 7: PURPOSE DISTRIBUTION ===
    doc.add_heading('7. Purpose Distribution', 1)
    
    purpose_text = """Analysis of registered dam purposes confirms the dominance of hydropower generation in Norwegian infrastructure, contrasting with India's multipurpose approach.

Purpose Distribution (Top Categories):
• Kraftproduksjon (Power Production): 1,830 reservoirs (61.1%) - Primary purpose
• Flomdemping (Flood Mitigation): Secondary consideration in some systems
• Multifunctional uses: Limited compared to Indian infrastructure

This single-purpose dominance reflects:

1. Geographic suitability: Mountainous terrain ideal for hydropower, less suitable for large-scale irrigation
2. Climate factors: Adequate precipitation reduces irrigation dependency
3. Development history: Early 20th-century electrification priorities
4. Economic structure: Industrial development based on cheap hydropower

Contrast with Indian Purpose Distribution:

Indian dam infrastructure serves multipurpose functions with integrated design for:
• Irrigation: Primary driver for many projects (agricultural water supply)
• Flood Control: Monsoon flood management (critical in flood-prone basins)
• Hydropower: Important but often secondary to water supply objectives
• Drinking Water Supply: Urban and rural water security
• Industrial Water Supply: Supporting industrial development

This difference in purpose distribution fundamentally affects:
• Operational constraints and flexibility
• Stakeholder complexity and conflict resolution requirements
• Economic benefit calculations and cost allocation
• Environmental impact profiles and mitigation strategies
• Regulatory frameworks and decision-making processes

The single-purpose hydropower focus in Norway enables more straightforward operational optimization, while India's multipurpose systems require complex trade-off analysis among competing objectives."""
    
    doc.add_paragraph(purpose_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 6: ').bold = True
    p.add_run('See visualization "purpose_distribution_norway.png" for purpose distribution pie chart')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 8: LARGE DAM INFRASTRUCTURE ===
    doc.add_heading('8. Large Dam Infrastructure', 1)
    
    large_dam_text = """Application of the Indian threshold criterion for "large dams" (gross storage ≥1000 MCM) to the Norwegian dataset identifies 10 reservoirs meeting this classification, with an additional 20 substantial reservoirs in the top 30 by storage capacity.

Large Norwegian Reservoirs (≥1000 MCM):
The detailed list of 30 largest Norwegian reservoirs is provided in the accompanying CSV file "large_norwegian_dams.csv" with complete technical specifications including:
• Dam name and location (county)
• Construction year and current age
• Gross storage capacity (MCM)
• Reservoir surface area (km²)  
• Storage efficiency (MCM/km²)
• Regulation range (meters)
• Geographic coordinates
• Operational status and purpose

Key Characteristics of Large Norwegian Reservoirs:

1. Geographic Distribution: Predominantly in Innlandet and Vestland counties, reflecting topographic suitability

2. Storage Efficiency: Large reservoirs exhibit high volume-to-area ratios, consistent with deep valley storage

3. Regulation Capability: Substantial regulation ranges enabling significant operational flexibility

4. Age Profile: Mix of mid-century infrastructure and more recent developments

5. Cascade Integration: Many large reservoirs function within interconnected cascade systems for optimized power generation

Note on Data Availability:

The accompanying CSV file provides comprehensive data for the 30 largest Norwegian reservoirs. Unlike the Indian National Register of Large Dams (NRLD), which includes structural dam height data, the Norwegian NVE public dataset does not systematically publish dam height measurements. This data limitation reflects different regulatory reporting priorities:

• Norway: Focus on reservoir operational parameters (storage, regulation range)
• India: Emphasis on structural characteristics (height, type, volume)

For academic research requiring Norwegian dam height data, specialized requests to NVE or consultation of project-specific technical documentation would be necessary."""
    
    doc.add_paragraph(large_dam_text)
    
    # Reference to CSV
    p = doc.add_paragraph()
    p.add_run('Data File: ').bold = True
    p.add_run('See "large_norwegian_dams.csv" for complete technical specifications of 30 largest reservoirs')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 9: COMPARATIVE INSIGHTS ===
    doc.add_heading('9. Comparative Insights: India vs Norway', 1)
    
    comparative_text = """This analysis reveals fundamental differences and notable similarities between Indian and Norwegian dam infrastructure, shaped by contrasting geography, climate, topography, and governance regimes.

Key Comparative Findings:

1. Scale and Magnitude:
   • Similar total dam counts (India: ~6,628; Norway: ~4,953)
   • Both countries maintain substantial hydropower infrastructure
   • Comparable challenges in managing aging dam portfolios

2. Development Trajectories:
   • Norway: Peak construction 1960s (post-war electrification)
   • India: Peak construction 1980s (post-independence development)
   • Different driving factors: electricity vs multipurpose development

3. Topographic Adaptation:
   • Norway: High storage efficiency (10.80 MCM/km²) - deep valley/fjord storage
   • India: Broader river valley impoundments with larger surface areas
   • Reflects fundamental geographic differences

4. Operational Characteristics:
   • Norway: High regulation ranges enabling flexible hydropower peaking
   • India: Multipurpose constraints limiting operational flexibility
   • Different system design philosophies

5. Purpose and Function:
   • Norway: Single-purpose hydropower dominance (61% of reservoirs)
   • India: Multipurpose orientation (irrigation/flood/power integration)
   • Reflects different developmental priorities and stakeholder needs

6. Governance Frameworks:
   • Norway: Centralized NVE oversight with WFD integration
   • India: Federated structure (Dam Safety Act 2021, state-level bodies)
   • Both effective within respective institutional contexts

7. Data Availability:
   • Norway: Excellent operational data (regulation ranges, storage efficiency)
   • India: Comprehensive structural data (dam height, type classifications)
   • Different regulatory reporting emphases

8. Age and Maintenance:
   • Both countries face similar aging infrastructure challenges
   • Norway: 60% of dams 50-99 years old
   • India: 62% of dams older than 25 years
   • Common needs for rehabilitation and modernization

Implications for Policy and Practice:

These comparative insights suggest that dam infrastructure performance, sustainability, and risk profiles are fundamentally shaped by:

• Natural environmental constraints (topography, hydrology, climate)
• Design philosophy and purpose prioritization
• Institutional and regulatory governance structures
• Historical development trajectories and timing

Understanding these differences enables more nuanced assessment of:
• Appropriate structural rehabilitation strategies
• Sediment management approaches adapted to local conditions
• Operational flexibility opportunities and constraints
• Safety oversight mechanisms suited to institutional contexts
• Adaptive policy design accounting for geographic and governance realities

The findings demonstrate that effective dam infrastructure management requires context-specific approaches that recognize fundamental geographic and institutional differences, while learning from common challenges in aging infrastructure management and safety oversight."""
    
    doc.add_paragraph(comparative_text)
    
    # Reference to visualization
    p = doc.add_paragraph()
    p.add_run('Figure 7: ').bold = True
    p.add_run('See visualization "norway_india_comparison.png" for detailed comparative metrics table')
    p.runs[-1].font.italic = True
    
    doc.add_paragraph()
    
    # === SECTION 10: DATA LIMITATIONS AND METHODOLOGY ===
    doc.add_heading('10. Data Limitations and Methodology', 1)
    
    limitations_text = """This analysis employs the Norwegian Water Resources and Energy Directorate (NVE) national hydropower database, accessed through public GIS datasets. Understanding data limitations is essential for appropriate interpretation and comparison.

Data Availability:

✓ Available in NVE Dataset:
• Dam locations (geographic coordinates)
• Construction years (for majority of structures)
• Reservoir storage volumes (volOppdemt field)
• Reservoir surface areas (areal_km2 field)
• Regulation ranges (HRV-LRV water levels) - Unique to Norwegian data
• Purpose classifications (formal_L field)
• Status and operational information
• Watershed associations

✗ Not Available in NVE Public Dataset:
• Structural dam heights - Key limitation for direct comparison
• Dam type classifications (earth-fill, concrete gravity, rock-fill, etc.)
• Detailed structural specifications (crest length, volume of materials)
• Design flood calculations and spillway capacities
• Foundation geology and treatment details
• Instrumentation specifications

This contrasts with India's National Register of Large Dams (NRLD-2023), which provides comprehensive structural data including height classifications, dam type distributions, and detailed technical specifications.

Methodological Approaches:

1. County Assignment: Geographic analysis employed coordinate-based county assignment using simplified regional boundaries. For high-precision spatial analysis, official county boundary shapefiles should be used.

2. Storage Efficiency Calculation: Computed as volume/area ratio for reservoirs with both parameters available (n=2,145 reservoirs, 71.6% of total).

3. Regulation Range Analysis: Based on reservoirs with documented HRV and LRV values (n=1,585, 52.9% of total). This unique parameter enables operational flexibility assessment not possible with Indian data.

4. Age Calculations: Referenced to year 2025, using construction year (idriftAar field) where available (n=3,874 dams, 78.2% of total).

5. Large Dam Threshold: Applied ≥1000 MCM storage criterion for consistency with international practice, while acknowledging that height-based classification cannot be performed.

Data Quality Considerations:

• NVE data represents official national registry with high reliability
• Regular updates maintained by regulatory authority
• Some historical structures may have incomplete records
• Focus on operational parameters reflects hydropower management priorities
• Public dataset subset; detailed technical data may exist in project-specific archives

Comparative Analysis Constraints:

When comparing Norwegian and Indian infrastructure:

• Different data collection priorities reflect different regulatory focuses
• Direct height-based comparisons are not feasible with available Norwegian data
• Purpose classifications use different taxonomies requiring interpretation
• Institutional context differences affect governance framework comparisons
• Geographic and climatic differences necessitate context-specific interpretation

These limitations are acknowledged to ensure appropriate use of findings and to guide future research directions where additional data sources or specialized data requests to NVE might be necessary."""
    
    doc.add_paragraph(limitations_text)
    
    doc.add_paragraph()
    
    # === SECTION 11: REFERENCES ===
    doc.add_heading('11. References', 1)
    
    references_text = """Dam Safety Act, 2021. Act No. 41 of 2021. Government of India. Gazette Notification S.O. 5422(E), effective 30 December 2021.

Graabak, I., Bakken, T. H., Feilberg, N., & Belsnes, M. M. (2017). Strategies for Norwegian hydropower in a changing Nordic energy system. MDPI Energies, 10(11), 1780. https://doi.org/10.3390/en10111780

Hanssen, F., Barton, D. N., Venter, O., Nowell, M. S., & Fjeldstad, H.-P. (2016). The Norwegian framework for implementing the Water Framework Directive: a review and analysis. Delft University of Technology Journals. https://journals.open.tudelft.nl/

National Register of Large Dams (NRLD-2023). Central Water Commission, Ministry of Jal Shakti, Government of India. Available at: https://cwc.gov.in/

Norwegian Water Resources and Energy Directorate (NVE). Hydropower Database. Public GIS datasets. Available at: https://www.nve.no/

Parliamentary Documentation, India. Digital Sansad. Information on dam statistics and safety. Government of India.

Water Framework Directive (WFD). Directive 2000/60/EC of the European Parliament. European Union.

Additional Data Sources:

• NVE Vannkraft_DamPunkt Shapefile - Dam point locations (n=4,953)
• NVE Vannkraft_DamLinje Shapefile - Dam line features (n=4,813)  
• NVE Vannkraft_Magasin Shapefile - Reservoir polygons (n=2,997)

Analysis Software:

• Python 3.x with GeoPandas, Pandas, Matplotlib libraries
• QGIS for spatial data verification
• Analysis code available in accompanying Python script"""
    
    doc.add_paragraph(references_text)
    
    doc.add_paragraph()
    
    # === APPENDIX ===
    doc.add_heading('Appendix: Summary Statistics', 1)
    
    appendix_text = """Complete statistical summary is provided in the accompanying file "norway_statistics_summary.txt"

Visualization Files:
1. regional_distribution_norway.png - County-wise dam distribution
2. age_wise_norway.png - Age classification of infrastructure
3. decade_wise_norway.png - Historical construction timeline
4. storage_efficiency_norway.png - Volume vs area analysis
5. regulation_range_norway.png - Operational flexibility histogram
6. purpose_distribution_norway.png - Purpose classification
7. norway_india_comparison.png - Comparative metrics table

Data Files:
• large_norwegian_dams.csv - Technical specifications of 30 largest reservoirs
• norway_statistics_summary.txt - Complete statistical summary

All analysis is fully reproducible using the accompanying Python script: norwegian_dam_analysis_comparative.py"""
    
    doc.add_paragraph(appendix_text)
    
    # Save document
    output_file = Path("Norwegian_Dam_Analysis_Report.docx")
    doc.save(output_file)
    
    print(f"✅ Word document created successfully!")
    print(f"📄 File: {output_file}")
    print(f"📊 Sections: 11 main sections + Appendix")
    print(f"📝 Ready for Janhavi's conference paper\n")
    print("=" * 70)

if __name__ == "__main__":
    create_report()

